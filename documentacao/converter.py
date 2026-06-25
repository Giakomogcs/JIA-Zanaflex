#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor Markdown -> Word (.docx) com identidade visual SENAI.
Gera o Playbook/Documentacao da Zanaflex a partir de DOCUMENTACAO.md.

Recursos:
- Capa estilizada (faixa vermelha SENAI)
- Cabecalho/rodape com numero de pagina (Pagina X de Y)
- Sumario (TOC) atualizavel do Word
- Titulos coloridos (H1 vermelho, H2 com regua, H3/H4)
- Tabelas zebradas com cabecalho vermelho
- Callouts NOTE/TIP/WARNING/DANGER com barra lateral e fundo tingido
- Passos numerados com badge
- Blocos de codigo com fundo escuro
- Markdown inline: **negrito**, *italico*, `codigo`
- Diagramas Mermaid renderizados em PNG via mmdc (fallback: bloco de codigo)
"""

import os
import re
import sys
import shutil
import subprocess
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- paleta SENAI
ACCENT   = "E30613"   # vermelho SENAI
ACCENT2  = "8B0410"   # vermelho escuro
INK      = "0B0B0B"   # titulos escuros
INK2     = "1F2937"   # texto corpo
GRAY     = "6B7280"   # legendas
RULE     = "D1D5DB"   # bordas
ZEBRA    = "F9F7F7"   # linha alternada
CELLHDR  = "FBE9EB"   # wash do cabecalho
CODE_BG  = "0B0B0B"
CODE_FG  = "E5E7EB"
WHITE    = "FFFFFF"

CALLOUT = {
    "NOTE":    {"bar": "1D4ED8", "bg": "EEF2FF", "title": "NOTA"},
    "INFO":    {"bar": "1D4ED8", "bg": "EEF2FF", "title": "INFO"},
    "TIP":     {"bar": "047857", "bg": "ECFDF5", "title": "DICA"},
    "WARNING": {"bar": "B45309", "bg": "FFFBEB", "title": "ATENCAO"},
    "DANGER":  {"bar": "B91C1C", "bg": "FEF2F2", "title": "CUIDADO"},
}

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE_DIR, "DOCUMENTACAO.md")
OUT_PATH = os.path.join(BASE_DIR, "Documentacao-Zanaflex.docx")
IMG_DIR = os.path.join(BASE_DIR, "_diagramas")


# --------------------------------------------------------------- helpers xml
def _set_shading(el, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    el.append(shd)


def _set_cell_bg(cell, color_hex):
    _set_shading(cell._tc.get_or_add_tcPr(), color_hex)


def _set_cell_borders(cell, color=RULE, sz=4, sides=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for side in sides:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)


def _no_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tcPr.append(borders)


def _set_left_bar(cell, color, sz=24):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)
    borders.append(e)
    tcPr.append(borders)


def _add_field(paragraph, instr):
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldBegin)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    run._r.append(instrText)
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldSep)
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldEnd)
    return run


# --------------------------------------------------------- inline formatting
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_inline(paragraph, text, base_color=INK2, base_size=10.5, base_bold=False):
    text = text.replace("\\n", "\n")
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = FONT_MONO
            r.font.color.rgb = RGBColor.from_string(ACCENT2)
            r.font.size = Pt(base_size - 0.5)
        else:
            r = paragraph.add_run(part)
            r.bold = base_bold
        r.font.color.rgb = RGBColor.from_string(base_color if not (part.startswith("`")) else ACCENT2)
        if not part.startswith("`"):
            r.font.name = FONT_BODY
            r.font.size = Pt(base_size)
            r.font.color.rgb = RGBColor.from_string(base_color)


# --------------------------------------------------------------- mermaid
def render_mermaid(code, idx):
    mmdc = shutil.which("mmdc")
    if not mmdc:
        return None
    os.makedirs(IMG_DIR, exist_ok=True)
    mmd = os.path.join(IMG_DIR, f"diag_{idx}.mmd")
    png = os.path.join(IMG_DIR, f"diag_{idx}.png")
    cfg = os.path.join(IMG_DIR, "puppeteer.json")
    with open(cfg, "w", encoding="utf-8") as f:
        f.write('{"args":["--no-sandbox","--disable-setuid-sandbox"]}')
    with open(mmd, "w", encoding="utf-8") as f:
        f.write(code)
    try:
        subprocess.run(
            [mmdc, "-i", mmd, "-o", png, "-b", "white", "-s", "2",
             "-p", cfg, "-t", "neutral"],
            check=True, capture_output=True, timeout=120, shell=False,
        )
        return png if os.path.exists(png) else None
    except Exception as e:
        sys.stderr.write(f"[mermaid] falha ao renderizar diag_{idx}: {e}\n")
        return None


# --------------------------------------------------------------- doc styling
def setup_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK2)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(19)
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        r.font.name = FONT_BODY
        _bottom_border(p, ACCENT, 12)
    elif level == 2:
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor.from_string(INK)
        r.font.name = FONT_BODY
        _bottom_border(p, RULE, 6)
    elif level == 3:
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = RGBColor.from_string(INK2)
        r.font.name = FONT_BODY
    else:
        pf.space_before = Pt(8)
        pf.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(GRAY)
        r.font.name = FONT_BODY
    return p


def _bottom_border(paragraph, color, sz):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# --------------------------------------------------------------- blocks
def add_code_block(doc, lines, lang=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, CODE_BG)
    _no_cell_borders(cell)
    cell.paragraphs[0].text = ""
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else " ")
        r.font.name = FONT_MONO
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(CODE_FG)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout(doc, kind, lines):
    spec = CALLOUT.get(kind, CALLOUT["NOTE"])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, spec["bg"])
    _no_cell_borders(cell)
    _set_left_bar(cell, spec["bar"], 28)
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(2)
    tr = title_p.add_run(spec["title"])
    tr.bold = True
    tr.font.size = Pt(10)
    tr.font.name = FONT_BODY
    tr.font.color.rgb = RGBColor.from_string(spec["bar"])
    body = " ".join(lines).strip()
    bp = cell.add_paragraph()
    bp.paragraph_format.space_after = Pt(0)
    add_inline(bp, body, base_color=INK2, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, header, rows):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(header):
        _set_cell_bg(hdr[i], ACCENT)
        _set_cell_borders(hdr[i], RULE, 4)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(htext.strip())
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = FONT_BODY
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        bg = ZEBRA if ri % 2 == 0 else WHITE
        for ci, ctext in enumerate(row):
            if ci >= len(cells):
                continue
            _set_cell_bg(cells[ci], bg)
            _set_cell_borders(cells[ci], RULE, 4)
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            add_inline(p, ctext.strip(), base_color=INK2, base_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_numbered_step(doc, num, text):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    badge = tbl.cell(0, 0)
    badge.width = Inches(0.35)
    _set_cell_bg(badge, ACCENT)
    _no_cell_borders(badge)
    bp = badge.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_after = Pt(0)
    br = bp.add_run(str(num))
    br.bold = True
    br.font.color.rgb = RGBColor.from_string(WHITE)
    br.font.size = Pt(10)
    br.font.name = FONT_BODY
    body = tbl.cell(0, 1)
    _no_cell_borders(body)
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, base_color=INK2, base_size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text, base_color=INK2, base_size=10.5)


def add_image_centered(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # largura maxima ~6 polegadas
    try:
        run.add_picture(path, width=Inches(6.0))
    except Exception:
        run.add_picture(path)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# --------------------------------------------------------------- cover/header
def add_cover(doc, title, subtitle, meta_lines):
    # faixa superior vermelha
    band = doc.add_table(rows=1, cols=1)
    c = band.cell(0, 0)
    _set_cell_bg(c, ACCENT)
    _no_cell_borders(c)
    for _ in range(2):
        c.add_paragraph()
    tp = c.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = tp.add_run("SENAI")
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor.from_string(WHITE)
    tr.font.name = FONT_BODY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    r.font.name = FONT_BODY
    _bottom_border(p, ACCENT2, 18)

    sp = doc.add_paragraph()
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor.from_string(INK2)
    sr.font.name = FONT_BODY

    for _ in range(8):
        doc.add_paragraph()

    for label, value in meta_lines:
        mp = doc.add_paragraph()
        lr = mp.add_run(f"{label}: ")
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = RGBColor.from_string(GRAY)
        lr.font.name = FONT_BODY
        vr = mp.add_run(value)
        vr.font.size = Pt(11)
        vr.font.color.rgb = RGBColor.from_string(INK2)
        vr.font.name = FONT_BODY

    # rodape vermelho da capa
    for _ in range(3):
        doc.add_paragraph()
    foot = doc.add_table(rows=1, cols=1)
    fc = foot.cell(0, 0)
    _set_cell_bg(fc, ACCENT2)
    _no_cell_borders(fc)
    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Documento de Transferência de Tecnologia — Confidencial")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string(WHITE)
    fr.font.name = FONT_BODY

    doc.add_page_break()


def add_toc(doc):
    h = doc.add_paragraph()
    r = h.add_run("SUMÁRIO")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    r.font.name = FONT_BODY
    _bottom_border(h, ACCENT, 8)
    p = doc.add_paragraph()
    _add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    nr = note.add_run("(Clique com o botão direito → Atualizar campo para gerar o índice.)")
    nr.italic = True
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def setup_header_footer(doc, doc_title, project):
    section = doc.sections[-1]
    section.different_first_page_header_footer = True

    header = section.header
    htbl = header.add_table(rows=1, cols=2, width=Inches(6.5))
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc = htbl.cell(0, 0).paragraphs[0]
    lr = lc.add_run(doc_title)
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor.from_string(GRAY)
    lr.font.name = FONT_BODY
    rc = htbl.cell(0, 1).paragraphs[0]
    rc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rc.add_run(project)
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor.from_string(ACCENT)
    rr.bold = True
    rr.font.name = FONT_BODY

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Página ")
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor.from_string(GRAY)
    _add_field(fp, "PAGE")
    r2 = fp.add_run(" de ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor.from_string(GRAY)
    _add_field(fp, "NUMPAGES")


# --------------------------------------------------------------- parser
def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def parse_and_build(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    diag_idx = 0

    # pula o bloco de titulo inicial (ja vai na capa)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # consome fechamento
            if lang == "mermaid":
                diag_idx += 1
                png = render_mermaid("\n".join(block), diag_idx)
                if png:
                    add_image_centered(doc, png, f"Diagrama {diag_idx}")
                else:
                    add_code_block(doc, block, "mermaid")
                    cap = doc.add_paragraph()
                    cr = cap.add_run("(diagrama Mermaid)")
                    cr.italic = True
                    cr.font.size = Pt(8.5)
                    cr.font.color.rgb = RGBColor.from_string(GRAY)
            else:
                add_code_block(doc, block, lang)
            continue

        # callout
        m = re.match(r"^>\s*\[!(\w+)\]\s*$", stripped)
        if m:
            kind = m.group(1).upper()
            i += 1
            body = []
            while i < n and lines[i].strip().startswith(">"):
                body.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(doc, kind, body)
            continue

        # heading
        hm = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+$", stripped):
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_separator_row(lines[i + 1]):
            header = split_table_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                if is_separator_row(lines[i]):
                    i += 1
                    continue
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # ordered list (numbered steps)
        om = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if om:
            add_numbered_step(doc, om.group(1), om.group(2))
            i += 1
            continue

        # bullet
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            add_bullet(doc, text)
            i += 1
            continue

        # blank
        if stripped == "":
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, stripped, base_color=INK2, base_size=10.5)
        i += 1


# --------------------------------------------------------------- main
def main():
    if not os.path.exists(MD_PATH):
        sys.exit(f"Markdown nao encontrado: {MD_PATH}")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md = f.read()

    # remove o cabecalho inicial (titulo + metadados) ate o primeiro '---'
    body = md
    parts = md.split("\n---\n", 1)
    if len(parts) == 2:
        body = parts[1]

    doc = Document()
    setup_base_styles(doc)

    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    add_cover(
        doc,
        "Documentação Zanaflex",
        "Agente de IA + RAG com ACL por equipe e integração com ERP",
        [
            ("Projeto", "Zanaflex — Agente IA / RAG"),
            ("Versão", "1.0"),
            ("Data", "24/06/2026"),
            ("Autor", "Doc Master"),
        ],
    )
    setup_header_footer(doc, "Documentação Zanaflex", "Zanaflex IA/RAG")
    add_toc(doc)
    parse_and_build(doc, body)

    doc.save(OUT_PATH)
    print(f"OK -> {OUT_PATH}")


if __name__ == "__main__":
    main()
